"""
Document extraction and analysis module.
Handles PDF, DOCX, text files, and images with built-in OCR support.
Unified tool that combines text extraction and Tesseract OCR.
"""

import io
import os
from typing import Optional, Tuple, List
from functools import lru_cache
import asyncio

# OCR processing
try:
    import pytesseract
    from PIL import Image

    HAS_PYTESSERACT = True
except ImportError:
    HAS_PYTESSERACT = False

try:
    from pdf2image import convert_from_bytes

    HAS_PDF2IMAGE = True
except ImportError:
    HAS_PDF2IMAGE = False

# PDF processing
try:
    import pdfplumber

    HAS_PDFPLUMBER = True
except ImportError:
    HAS_PDFPLUMBER = False

try:
    from pypdf import PdfReader

    HAS_PYPDF = True
except ImportError:
    HAS_PYPDF = False

# DOCX processing
try:
    from docx import Document as DocxDocument

    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


def check_ocr_availability() -> Tuple[bool, str]:
    """
    Check if OCR is available on the system.

    Returns:
        Tuple of (is_available, message)
    """
    if not HAS_PYTESSERACT:
        return False, "pytesseract not installed"

    try:
        version = pytesseract.get_tesseract_version()
        return True, f"Tesseract {version} available"
    except Exception as e:
        return False, f"Tesseract not found: {str(e)}"


class DocumentExtractor:
    """
    Unified document text extraction with built-in OCR support.
    Supports PDF, DOCX, plain text files, and images (via Tesseract OCR).
    """

    SUPPORTED_EXTENSIONS = {
        ".pdf",
        ".docx",
        ".doc",
        ".txt",
        ".jpg",
        ".jpeg",
        ".png",
        ".bmp",
        ".tiff",
        ".gif",
    }
    SUPPORTED_IMAGE_EXTENSIONS = {".jpg", ".jpeg", ".png", ".bmp", ".tiff", ".gif"}
    MAX_PAGES = 100  # Limit for performance
    MAX_IMAGE_SIZE = (3000, 3000)  # Resize large images for faster OCR

    def __init__(self, tesseract_cmd: Optional[str] = None):
        self._cache = {}
        self._ocr_available = False

        if HAS_PYTESSERACT:
            if tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
            is_available, message = check_ocr_availability()
            self._ocr_available = is_available
            if is_available:
                print(f"OCR enabled: {message}")
            else:
                print(f"OCR not available: {message}")

    # ========================================================================
    # Public API
    # ========================================================================

    async def extract_text(self, file_bytes: bytes, filename: str) -> Tuple[str, str]:
        """
        Extract text from a document file or image.

        Args:
            file_bytes: The file content as bytes
            filename: Original filename to determine type

        Returns:
            Tuple of (extracted_text, file_type)
        """
        extension = self._get_extension(filename)

        # Handle images with OCR
        if extension in self.SUPPORTED_IMAGE_EXTENSIONS:
            if self._ocr_available:
                text = await self._extract_from_image(file_bytes, filename)
                return text, "image_ocr"
            else:
                raise ValueError(
                    "OCR not available. Please install tesseract-ocr system package and pytesseract."
                )

        if extension == ".pdf":
            text = await self._extract_pdf_async(file_bytes)

            # If PDF has very little text, try OCR as fallback (likely scanned)
            if len(text.strip()) < 100 and self._ocr_available:
                try:
                    ocr_text = await self._extract_from_pdf_images(file_bytes)
                    if len(ocr_text.strip()) > len(text.strip()):
                        return ocr_text, "pdf_ocr"
                except Exception as e:
                    print(f"OCR fallback failed: {e}")

            return text, "pdf"
        elif extension in (".docx", ".doc"):
            text = await self._extract_docx_async(file_bytes)
            return text, "docx"
        elif extension == ".txt":
            text = file_bytes.decode("utf-8", errors="ignore")
            return text, "txt"
        else:
            raise ValueError(f"Unsupported file type: {extension}")

    def get_text_preview(self, text: str, max_length: int = 500) -> str:
        """Get a preview of the extracted text."""
        if len(text) <= max_length:
            return text
        return text[:max_length] + "..."

    def is_image_file(self, filename: str) -> bool:
        """Check if filename is a supported image type."""
        extension = os.path.splitext(filename.lower())[1]
        return extension in self.SUPPORTED_IMAGE_EXTENSIONS

    # ========================================================================
    # OCR Extraction
    # ========================================================================

    async def _extract_from_image(self, image_bytes: bytes, filename: str = "") -> str:
        """
        Extract text from an image file using OCR.

        Args:
            image_bytes: Image file bytes
            filename: Original filename (for logging)

        Returns:
            Extracted text
        """
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self._extract_from_image_sync, image_bytes, filename
        )

    def _extract_from_image_sync(self, image_bytes: bytes, filename: str = "") -> str:
        """Synchronous image OCR extraction."""
        try:
            image = Image.open(io.BytesIO(image_bytes))

            # Convert to RGB if needed
            if image.mode not in ("RGB", "L"):
                image = image.convert("RGB")

            # Resize large images for faster processing
            if (
                image.size[0] > self.MAX_IMAGE_SIZE[0]
                or image.size[1] > self.MAX_IMAGE_SIZE[1]
            ):
                image.thumbnail(self.MAX_IMAGE_SIZE, Image.Resampling.LANCZOS)

            text = pytesseract.image_to_string(image, lang="eng")
            return text.strip()

        except Exception as e:
            raise ValueError(f"Failed to extract text from image: {str(e)}")

    async def _extract_from_pdf_images(
        self, pdf_bytes: bytes, max_pages: int = 20
    ) -> str:
        """
        Extract text from scanned PDF by converting pages to images and OCR.

        Args:
            pdf_bytes: PDF file bytes
            max_pages: Maximum number of pages to process

        Returns:
            Extracted text from all pages
        """
        if not HAS_PDF2IMAGE:
            raise ImportError(
                "pdf2image not installed. Install with: pip install pdf2image"
            )

        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(
            None, self._extract_from_pdf_images_sync, pdf_bytes, max_pages
        )

    def _extract_from_pdf_images_sync(
        self, pdf_bytes: bytes, max_pages: int = 20
    ) -> str:
        """Synchronous PDF OCR extraction."""
        try:
            images = convert_from_bytes(
                pdf_bytes,
                first_page=1,
                last_page=max_pages,
                dpi=200,
                fmt="jpeg",
            )

            text_parts = []
            for i, image in enumerate(images, 1):
                if (
                    image.size[0] > self.MAX_IMAGE_SIZE[0]
                    or image.size[1] > self.MAX_IMAGE_SIZE[1]
                ):
                    image.thumbnail(self.MAX_IMAGE_SIZE, Image.Resampling.LANCZOS)

                page_text = pytesseract.image_to_string(image, lang="eng")
                if page_text.strip():
                    text_parts.append(f"--- Page {i} ---\n{page_text.strip()}")

            if not text_parts:
                return ""

            full_text = "\n\n".join(text_parts)

            if len(images) >= max_pages:
                full_text += f"\n\n[Note: Only first {max_pages} pages processed]"

            return full_text

        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF images: {str(e)}")

    # ========================================================================
    # PDF / DOCX Extraction
    # ========================================================================

    async def _extract_pdf_async(self, file_bytes: bytes) -> str:
        """Extract text from PDF asynchronously."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_pdf_sync, file_bytes)

    def _extract_pdf_sync(self, file_bytes: bytes) -> str:
        """
        Extract text from PDF synchronously.
        Uses pdfplumber for better text extraction, falls back to pypdf.
        """
        file_stream = io.BytesIO(file_bytes)

        # Try pdfplumber first (better for tables and complex layouts)
        if HAS_PDFPLUMBER:
            try:
                text_parts = []
                with pdfplumber.open(file_stream) as pdf:
                    for i, page in enumerate(pdf.pages):
                        if i >= self.MAX_PAGES:
                            text_parts.append(
                                f"\n[Truncated: Document has more than {self.MAX_PAGES} pages]"
                            )
                            break
                        page_text = page.extract_text()
                        if page_text:
                            text_parts.append(page_text)
                return "\n\n".join(text_parts)
            except Exception:
                file_stream.seek(0)  # Reset stream for fallback

        # Fallback to pypdf
        if HAS_PYPDF:
            try:
                reader = PdfReader(file_stream)
                text_parts = []
                for i, page in enumerate(reader.pages):
                    if i >= self.MAX_PAGES:
                        text_parts.append(
                            f"\n[Truncated: Document has more than {self.MAX_PAGES} pages]"
                        )
                        break
                    text_parts.append(page.extract_text() or "")
                return "\n\n".join(text_parts)
            except Exception as e:
                raise ValueError(f"Failed to extract PDF text: {str(e)}")

        raise ValueError(
            "No PDF processing library available. Install pdfplumber or pypdf."
        )

    async def _extract_docx_async(self, file_bytes: bytes) -> str:
        """Extract text from DOCX asynchronously."""
        loop = asyncio.get_event_loop()
        return await loop.run_in_executor(None, self._extract_docx_sync, file_bytes)

    def _extract_docx_sync(self, file_bytes: bytes) -> str:
        """Extract text from DOCX synchronously."""
        if not HAS_DOCX:
            raise ValueError(
                "python-docx not installed. Install it with: pip install python-docx"
            )

        file_stream = io.BytesIO(file_bytes)
        doc = DocxDocument(file_stream)

        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)

        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    paragraphs.append(row_text)

        return "\n\n".join(paragraphs)

    # ========================================================================
    # Helpers
    # ========================================================================

    def _get_extension(self, filename: str) -> str:
        """Get lowercase file extension."""
        return os.path.splitext(filename.lower())[1]


# Singleton instance
_extractor: Optional[DocumentExtractor] = None


def get_document_extractor() -> DocumentExtractor:
    """Get or create the document extractor instance."""
    global _extractor
    if _extractor is None:
        _extractor = DocumentExtractor()
    return _extractor
